from docx import Document

# home-made function to remove a paragraph
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# read file. Only docx files are available
document = Document('D:\\beloved\\20200412word_process\\2018unit2.docx')

paragraphs = document.paragraphs

for p in paragraphs:
    text = p.text
    if text.startswith('【正确答案】'):
        answer = text[-1]
    elif text.startswith('【答案解析】'):
        text += '因此本题选' + answer + '。'
        p.text = text
        run = p.add_run()
        run.add_break()
    elif text.startswith('【考点还原】'):
        delete_paragraph(p)
document.save('output.docx')
